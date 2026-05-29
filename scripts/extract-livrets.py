"""Extract text from all welcome books to analyze structure."""
import sys
from pathlib import Path
import pypdf
from docx import Document

DOWNLOADS = Path(r"C:\Users\Xensma\Downloads")
OUT = Path(r"C:\Users\Xensma\corail-livret\scripts\livrets-extracted")
OUT.mkdir(exist_ok=True)

FILES = [
    ("Corail Conciergerie Guide de Bienvenue Crystal BEACH.pdf", "crystal-beach"),
    ("Livret d'accueil Villa Piment Bleu.pdf", "piment-bleu"),
    ("Livret d'accueil Villa Rêve Indigo.pdf", "reve-indigo"),
    ("Livret d'accueil Villa Tiki Reva.pdf", "tiki-reva"),
    ("Savannah.pdf", "savannah"),
    ("Livret Accueil Le Dattier Villa _ corail.docx", "dattier"),
    ("Livret Accueil villa La Rhumière _ corail.docx", "rhumiere"),
]

def extract_pdf(path: Path) -> str:
    reader = pypdf.PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text()
        except Exception as e:
            text = f"[ERROR page {i+1}: {e}]"
        pages.append(f"\n=== PAGE {i+1} ===\n{text}")
    return "\n".join(pages)

def extract_docx(path: Path) -> str:
    doc = Document(str(path))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # also tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(f"[CELL] {cell.text}")
    return "\n".join(parts)

for filename, slug in FILES:
    src = DOWNLOADS / filename
    if not src.exists():
        print(f"MISSING: {filename}")
        continue
    print(f"Extracting {filename}...")
    if src.suffix.lower() == ".pdf":
        content = extract_pdf(src)
    else:
        content = extract_docx(src)
    (OUT / f"{slug}.txt").write_text(content, encoding="utf-8")
    print(f"  -> {slug}.txt ({len(content)} chars)")

print("\nAll done.")
